#!/usr/bin/env python3
# -*- coding: utf-8 -*-
'''
Author: Theo Portlock
'''
import functions as f
import numpy as np
import pandas as pd
from docx import Document

# for samplesheet
docs = []
docs.append(Document('../data/InoliWD_Batch3_JOS_CS-ANZ DNA Sample Information Form.docx'))
docs.append(Document('../data/InoliWD_Batch4_JOS_CS-ANZ DNA Sample Information Form.docx'))
docs.append(Document('../data/InoliWD_Batch5_JOS_CS-ANZ DNA Sample Information Form.docx'))
docs.append(Document('../data/InoliWD_JOS_CS-ANZ DNA Sample Information Form.docx'))
output = []
for i, doc in enumerate(docs):
    tables = doc.tables[0]
    coltext = []
    for rows in tables.rows[12:]:
        rowtext = []
        for cell in rows.cells:
            rowtext.append(cell.text)
        coltext.append(rowtext)
    df = pd.DataFrame(coltext)
    df = df.T.drop_duplicates().T
    df = df.iloc[:,:-1] # remove remark column
    df.columns = df.iloc[0]
    df = df.iloc[1:,:]
    df = df.set_index('*Name on tube')
    df = df.drop_duplicates()
    output.append(df)
df = pd.concat(output).sort_index().drop_duplicates().iloc[1:]
df = df['Name on report'].drop_duplicates().to_frame().reset_index().set_index('Name on report').drop_duplicates() # remove that duplicate - can use this in future to look at batch effect
df = df.loc[df.index.str.startswith('LCC')] # remove controls
df.loc[(df['*Name on tube'].str[:2] == 'C1') | (df['*Name on tube'].str[:2] == 'C2'), 'batch'] = 1
df.loc[df['*Name on tube'].str[:2] == 'C5', 'batch'] = 2
df.loc[df['*Name on tube'].str[:1] == 'J', 'batch'] = 3
df.loc[df['*Name on tube'].str[:1] == 'M', 'batch'] = 4
df.loc[df['*Name on tube'].str[:1] == 'S', 'batch'] = 5
df['ID'] = df.index.str[:7]
df.loc[df.index.str[7:] == '1001', 'timepoint'] = '000'
df.loc[df.index.str[3] == '3', 'timepoint'] = '104'
df.loc[df.index.str[7:] == '1002', 'timepoint'] = '012'
df.loc[df.index.str.contains(r'2\d\d\d1002'), 'timepoint'] = '052'
df.loc[df.index.str[7:] == '1003', 'timepoint'] = '052'
samplesheet = df.dropna()
samplesheet = samplesheet.rename(columns={'*Name on tube':'Seq_ID'})
samplesheet.index = samplesheet.ID + samplesheet.timepoint
samplesheet.index.name = "TimeID"
#f.save(samplesheet, 'samplesheet')

# for quality
quality = pd.read_csv('../data/m4efad_kneaddata_read_counts_oct2023.tsv', sep='\t', index_col=0)
quality = quality.loc[:,quality.columns.str.contains('final')].sum(axis=1).div(1e6).to_frame('HQ read depth (million reads)')
quality = quality.join(samplesheet.reset_index().set_index('Seq_ID')['TimeID'], how='inner').set_index('TimeID')
# Remove nonbaseline data
quality = quality.loc[quality.index.str[3] !='3']
quality = quality.loc[quality.index.str[-1] !='2']
quality.index = quality.index.str[:-3]
quality.index = quality.index.rename('ID')
f.save(quality, 'quality')

# for taxonomy
taxo = pd.read_csv('../data/m4efad_metaphlan3_profiles_oct2023.tsv', sep='\t', index_col=0, header=1)
taxo.columns = taxo.columns.str.replace('\.metaphlan','', regex=True)
taxo = taxo.T.join(samplesheet.reset_index().set_index('Seq_ID')['TimeID'], how='inner').set_index('TimeID')
taxo = taxo.loc[taxo.index.str[3] !='3']
taxo = taxo.loc[taxo.index.str[-1] !='2']
taxo.index = taxo.index.str[:-3]
taxo.index = taxo.index.rename('ID')
taxo.columns = taxo.columns.str.replace('.*\|','',regex=True)
f.save(taxo, 'taxo')

# for functions - in cpm
pathways = pd.read_csv("../data/m4efad_humann3_pathway_cpm_oct2023.tsv", index_col=0, sep='\t').T
pathways.index = pathways.index.str.replace('_Abundance','')
pathways = pathways.join(samplesheet.reset_index().set_index('Seq_ID')['TimeID'], how='inner').set_index('TimeID')
pathways = pathways.loc[pathways.index.str[3] !='3']
pathways = pathways.loc[pathways.index.str[-1] !='2']
pathways.index = pathways.index.str[:-3]
pathways = f.filter(pathways, nonzero=True)
pathways.index = pathways.index.rename('ID')
f.save(pathways, 'pathwaysstrat')
pathwaysall = f.norm(pathways.loc[:, ~pathways.columns.str.contains('\|')])
f.save(pathwaysall, 'pathways')

# for melonnpan
melon = pd.read_csv("../data/melonnpan_results2MelonnPan_Predicted_Metabolites.txt", index_col=0, sep='\t')
melon = melon.join(samplesheet.reset_index().set_index('Seq_ID')['TimeID'], how='inner').set_index('TimeID')
melon = melon.loc[melon.index.str[3] !='3']
melon = melon.loc[melon.index.str[-1] !='2']
melon.index = melon.index.str[:-3]
melon = f.filter(melon, nonzero=True)
melon.index = melon.index.rename('ID')
f.save(melon, 'melon')
